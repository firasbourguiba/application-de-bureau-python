"""
Module pour générer le document Word final : page de titre (titre du livre,
image fusionnée, auteur, auteur du rapport) et page de graphique (distribution
des paragraphes + description + statistiques).
"""

import logging
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class WordGenerationError(Exception):
    """Exception levée en cas d'erreur lors de la génération du document Word."""
    pass


class WordReportGenerator:
    """Génère le rapport Word final à partir des éléments produits par l'analyse."""

    def __init__(self):
        self.doc: Document | None = None

    def create_document(self) -> Document:
        """Crée un nouveau document Word vide."""
        self.doc = Document()
        return self.doc

    def add_title_page(self, book_title: str, book_author: str,
                        report_author: str, image_path: str = None) -> None:
        """
        Ajoute la page de titre : titre du livre, image n°1 (avec logo fusionné),
        auteur du livre, et auteur du rapport.

        Args:
            book_title: Titre du livre.
            book_author: Auteur du livre.
            report_author: Nom de l'auteur du rapport (l'étudiant).
            image_path: Chemin vers l'image fusionnée (image n°1 + logo).

        Raises:
            WordGenerationError: en cas d'erreur d'insertion.
        """
        if self.doc is None:
            self.create_document()

        try:
            title_para = self.doc.add_heading(book_title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_para.runs[0]
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

            self.doc.add_paragraph()

            if image_path and Path(image_path).exists():
                img_para = self.doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_para.add_run().add_picture(image_path, width=Inches(4.0))
            else:
                logger.warning(f"Image introuvable pour la page de titre : {image_path}")

            self.doc.add_paragraph()
            self.doc.add_paragraph()

            author_para = self.doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(f"Auteur du livre : {book_author}")
            author_run.font.size = Pt(14)
            author_run.font.italic = True

            self.doc.add_paragraph()

            report_para = self.doc.add_paragraph()
            report_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            report_run = report_para.add_run(f"Rapport réalisé par : {report_author}")
            report_run.font.size = Pt(13)
            report_run.font.bold = True

            self.doc.add_page_break()

        except Exception as e:
            logger.error(f"Erreur lors de la création de la page de titre : {e}")
            raise WordGenerationError(f"Erreur lors de la création de la page de titre : {e}")

        logger.info("Page de titre ajoutée")

    def add_graph_page(self, chart_path: str, description: str, statistics: dict,
                        source_text: str) -> None:
        """
        Ajoute la page de graphique : graphique de distribution, description
        textuelle, et tableau de statistiques détaillées.

        Args:
            chart_path: Chemin vers l'image PNG du graphique.
            description: Texte décrivant l'intrigue / le contenu analysé.
            statistics: Dictionnaire de statistiques (total_paragraphs, total_words,
                        min_words, max_words, avg_words).
            source_text: Texte décrivant la source des données.

        Raises:
            WordGenerationError: en cas d'erreur d'insertion.
        """
        if self.doc is None:
            self.create_document()

        try:
            heading = self.doc.add_heading("Analyse du premier chapitre", level=1)
            heading.runs[0].font.bold = True
            heading.runs[0].font.size = Pt(20)
            heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

            if chart_path and Path(chart_path).exists():
                chart_para = self.doc.add_paragraph()
                chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chart_para.add_run().add_picture(chart_path, width=Inches(6.0))
            else:
                logger.warning(f"Graphique introuvable : {chart_path}")

            self.doc.add_paragraph()

            desc_heading = self.doc.add_heading("Description", level=2)
            desc_heading.runs[0].font.bold = True
            desc_heading.runs[0].font.italic = True
            desc_heading.runs[0].font.size = Pt(14)

            desc_para = self.doc.add_paragraph(description)
            desc_para.runs[0].font.size = Pt(11)

            self.doc.add_paragraph()

            stats_heading = self.doc.add_heading("Statistiques", level=2)
            stats_heading.runs[0].font.bold = True
            stats_heading.runs[0].font.italic = True
            stats_heading.runs[0].font.size = Pt(14)

            self._add_statistics_table(statistics)

            self.doc.add_paragraph()
            source_para = self.doc.add_paragraph()
            source_run = source_para.add_run(f"Source des données : {source_text}")
            source_run.font.italic = True
            source_run.font.size = Pt(9)

        except Exception as e:
            logger.error(f"Erreur lors de la création de la page de graphique : {e}")
            raise WordGenerationError(f"Erreur lors de la création de la page de graphique : {e}")

        logger.info("Page de graphique ajoutée")

    def _add_statistics_table(self, statistics: dict) -> None:
        """Ajoute un tableau de statistiques formaté dans le document."""
        table = self.doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"

        header_cells = table.rows[0].cells
        header_cells[0].text = "Métrique"
        header_cells[1].text = "Valeur"
        for cell in header_cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True

        labels = [
            ("Nombre de paragraphes", "total_paragraphs"),
            ("Nombre total de mots", "total_words"),
            ("Minimum de mots par paragraphe", "min_words"),
            ("Maximum de mots par paragraphe", "max_words"),
            ("Moyenne de mots par paragraphe", "avg_words"),
        ]

        for label, key in labels:
            if key in statistics:
                row = table.add_row().cells
                row[0].text = label
                value = statistics[key]
                row[1].text = f"{value:.2f}" if isinstance(value, float) else str(value)

    def save_document(self, filepath: str) -> None:
        """
        Sauvegarde le document Word sur le disque.

        Args:
            filepath: Chemin de destination du fichier .docx.

        Raises:
            WordGenerationError: si aucun document n'existe ou en cas d'échec d'écriture.
        """
        if self.doc is None:
            raise WordGenerationError("Aucun document à sauvegarder.")

        try:
            Path(filepath).parent.mkdir(parents=True, exist_ok=True)
            self.doc.save(filepath)
        except Exception as e:
            logger.error(f"Erreur lors de la sauvegarde du document Word : {e}")
            raise WordGenerationError(f"Impossible de sauvegarder le document dans '{filepath}' : {e}")

        logger.info(f"Document Word sauvegardé dans {filepath}")
